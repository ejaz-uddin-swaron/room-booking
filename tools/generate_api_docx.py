from docx import Document
from docx.shared import Pt


def add_heading(document, text, level=1):
    document.add_heading(text, level=level)


def add_paragraph(document, text, bold=False):
    p = document.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    run.font.size = Pt(11)


def add_bullet(document, text):
    p = document.add_paragraph(text, style='List Bullet')
    p.runs[0].font.size = Pt(11)


def add_code(document, text):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)


def main():
    base_url = "https://room-booking-pjo6.onrender.com/api"

    doc = Document()
    doc.add_heading('VillaEase Backend API Documentation', 0)

    add_heading(doc, 'Overview', 2)
    add_paragraph(doc, 'This document describes the RESTful API for the VillaEase room booking platform. Responses are JSON and JWT is used for protected/admin operations.')

    add_heading(doc, 'Base', 2)
    add_paragraph(doc, f"Base URL: {base_url}")
    add_paragraph(doc, 'Headers:')
    add_bullet(doc, 'Content-Type: application/json')
    add_bullet(doc, 'Authorization: Bearer <jwt_token> (for protected routes)')
    add_paragraph(doc, 'Standard Response:')
    add_code(doc, '{\n  "success": true,\n  "data": { ... }\n}')
    add_paragraph(doc, 'Error Response:')
    add_code(doc, '{\n  "success": false,\n  "error": "Message",\n  "status": 400\n}')

    # Auth
    add_heading(doc, 'Authentication', 2)
    add_paragraph(doc, 'POST /auth/login')
    add_code(doc, '{\n  "username": "admin",\n  "password": "admin123"\n}')
    add_paragraph(doc, 'Response:')
    add_code(doc, '{\n  "success": true,\n  "data": {\n    "token": "<jwt>",\n    "refresh": "<token>",\n    "user": { "id": 1, "username": "admin", "role": "admin" }\n  }\n}')
    add_paragraph(doc, 'GET /auth/verify (JWT)')
    add_code(doc, '{\n  "success": true,\n  "data": {\n    "valid": true,\n    "user": { "id": 1, "username": "admin", "role": "admin" }\n  }\n}')
    add_paragraph(doc, 'POST /auth/logout (JWT)')
    add_code(doc, '{ "refresh": "<refresh_token>" }')
    add_paragraph(doc, 'POST /auth/jwt/refresh')
    add_code(doc, '{ "refresh": "<refresh_token>" } -> { "access": "<new_access>" }')

    # Rooms
    add_heading(doc, 'Rooms', 2)
    add_paragraph(doc, 'GET /rooms')
    add_paragraph(doc, 'Query params: location, check_in, check_out, guests, min_price, max_price, room_type, amenities (repeat), page, page_size')
    add_paragraph(doc, 'Response: array of rooms')
    add_paragraph(doc, 'GET /rooms/:id')
    add_paragraph(doc, 'Response: single room')
    add_paragraph(doc, 'POST /rooms (Admin only)')
    add_code(doc, '{ "name": "Luxury Ocean View Suite", "type": "Suite", "price": 250, "images": ["..."], "amenities": ["WiFi"], "description": "...", "location": "London", "maxGuests": 4, "bedrooms": 2, "bathrooms": 2, "size": 85, "available": true }')
    add_paragraph(doc, 'PUT /rooms/:id (Admin only) – same fields as create, all optional')
    add_paragraph(doc, 'DELETE /rooms/:id (Admin only)')

    # Bookings
    add_heading(doc, 'Bookings', 2)
    add_paragraph(doc, 'POST /bookings')
    add_code(doc, '{\n  "roomId": "<room_id>",\n  "checkIn": "2024-02-01",\n  "checkOut": "2024-02-05",\n  "guests": 2,\n  "guestInfo": { "name": "John Doe", "email": "john@example.com", "phone": "+1234567890" }\n}')
    add_paragraph(doc, 'Response includes totalPrice and status (pending by default)')
    add_paragraph(doc, 'GET /bookings (Admin only)')
    add_paragraph(doc, 'Returns list of bookings with embedded room { id, name, location }')
    add_paragraph(doc, 'PATCH /bookings/:id/status (Admin only)')
    add_code(doc, '{ "status": "confirmed" }  # pending|confirmed|cancelled|completed')

    # Uploads
    add_heading(doc, 'Uploads', 2)
    add_paragraph(doc, 'POST /upload/images (JWT)')
    add_paragraph(doc, 'multipart/form-data with images[]; jpg/jpeg/png/webp; max 5MB/file')
    add_paragraph(doc, 'Response returns absolute URLs')

    # Admin Stats
    add_heading(doc, 'Admin Stats', 2)
    add_paragraph(doc, 'GET /admin/stats (Admin only)')
    add_paragraph(doc, 'Response: { totalRooms, totalBookings, totalRevenue, occupancyRate }')

    # Errors & Notes
    add_heading(doc, 'Errors', 2)
    add_bullet(doc, '401 Unauthorized – missing/invalid token')
    add_bullet(doc, '403 Forbidden – insufficient privileges')
    add_bullet(doc, '404 Not Found – resource not found')
    add_bullet(doc, '422 Validation Error – invalid input')

    add_heading(doc, 'Notes', 2)
    add_bullet(doc, 'Dates format: YYYY-MM-DD; checkIn < checkOut; checkIn must be in the future')
    add_bullet(doc, 'Bookings prevent overlaps for pending/confirmed')
    add_bullet(doc, 'totalPrice = price per night × nights')
    add_bullet(doc, 'Rooms list supports pagination: page, page_size (default page_size = 50)')

    doc.save('VillaEase_API_Documentation.docx')


if __name__ == '__main__':
    main()
