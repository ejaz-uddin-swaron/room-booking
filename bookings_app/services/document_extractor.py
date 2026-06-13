import requests
import fitz  # PyMuPDF
import docx
from io import BytesIO

def extract_pdf_text(file_bytes):
    """Extract raw text from PDF bytes using PyMuPDF (fitz)."""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    except Exception as e:
        return f"[Error extracting PDF text: {str(e)}]"

def extract_docx_text(file_bytes):
    """Extract raw text from DOCX bytes using python-docx."""
    try:
        doc = docx.Document(BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs]
        
        # Parse tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                # Filter out duplicate empty cells
                if any(row_text):
                    table_text.append(" | ".join(row_text))
        
        full_text = "\n".join(paragraphs)
        if table_text:
            full_text += "\n\nTable Data:\n" + "\n".join(table_text)
        return full_text
    except Exception as e:
        return f"[Error extracting Word document text: {str(e)}]"

def extract_text_from_url(file_url, file_name=None):
    """Downloads a file and extracts text based on file type."""
    if not file_url:
        return ""
    
    try:
        response = requests.get(file_url, timeout=30)
        response.raise_for_status()
        file_bytes = response.content
    except Exception as e:
        return f"[Error fetching document from URL: {str(e)}]"
    
    # Determine extension
    name = file_name or file_url.split('/')[-1].split('?')[0]
    ext = name.split('.')[-1].lower() if '.' in name else ''
    
    if ext == 'pdf':
        return extract_pdf_text(file_bytes)
    elif ext in ['docx', 'doc']:
        return extract_docx_text(file_bytes)
    else:
        return f"[Unsupported file extension for extraction: {ext}]"
